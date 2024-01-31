import tkinter as tk
from tkinter import filedialog, ttk
from docx import Document
from deep_translator import GoogleTranslator


def selecionar_arquivo():
    caminho_arquivo = filedialog.askopenfilename(
        title="Selecione um arquivo Word", filetypes=[("Documentos Word", "*.docx")])
    if caminho_arquivo:
        traduzir_documento(caminho_arquivo)


def ajustar_nome(txt):
    return txt.replace(' ', '_').replace('.docx', '').lower()


def traduzir_documento(caminho_arquivo):
    documento = Document(caminho_arquivo)

    # Ask the user where to save the translated document
    destino_arquivo = filedialog.asksaveasfilename(
        defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")])

    if not destino_arquivo:
        label_status["text"] = "Tradução cancelada. Nenhum arquivo selecionado."
        return

    linguagem_destino = combo_linguagem_destino.get()

    for paragrafo in documento.paragraphs:
        texto_original = paragrafo.text
        texto_traduzido = GoogleTranslator(
            source='auto', target=linguagem_destino).translate(texto_original)
        paragrafo.clear()
        paragrafo.add_run(texto_traduzido)

    novo_nome = ajustar_nome(destino_arquivo)

    # Salvar documento docx
    nome_traduzido = f"{novo_nome}.docx"

    documento.save(nome_traduzido)

    label_status["text"] = f"Documento traduzido e salvo."


# Criar a interface gráfica
root = tk.Tk()
root.title("Tradutor Word")

# Define o tamanho da janela
largura = 300
altura = 200
x = (root.winfo_screenwidth() - largura) // 2
y = (root.winfo_screenheight() - altura) // 2
root.geometry(f"{largura}x{altura}+{x}+{y}")

# Label para exibir texto antes de selecionar arquivo
label_texto = tk.Label(
    root, text="Traduza seu Documento", font=("Helvetica", 14))
label_texto.pack(pady=10)

# Combobox para selecionar o idioma de destino
label_linguagem_destino = tk.Label(root, text="Selecione o Idioma de Destino:")
label_linguagem_destino.pack(pady=5)

# Lista de idiomas disponíveis
idiomas_disponiveis = ['pt', 'es', 'fr', 'de', 'it', 'ja', 'ko', 'zh-CN']

# Ajuste da largura e definição do Combobox
combo_linguagem_destino = ttk.Combobox(
    root, values=idiomas_disponiveis, state="readonly", width=5)
combo_linguagem_destino.set("pt")  # Valor padrão
combo_linguagem_destino.pack(pady=5)

# Botão para selecionar arquivo
btn_selecionar_arquivo = tk.Button(
    root, text="Selecionar Arquivo", command=selecionar_arquivo)
btn_selecionar_arquivo.pack(pady=20)

# Label para exibir status
label_status = tk.Label(root, text="")
label_status.pack()

# Iniciar o loop da interface gráfica
root.mainloop()
